"""文档解析模块

支持多种文档格式的解析，包括扫描件 PDF 的 OCR 处理。
"""

import os
import shutil
import subprocess
import tempfile
from pathlib import Path

from .exceptions import (
    EmptyFileError,
    MinerUNotInstalledError,
    MinerUTokenRequiredError,
    ParseError,
    UnsupportedFormatError,
)


def parse_file(filepath: str) -> str:
    """解析文档文件，返回纯文本

    支持: .md, .txt, .docx, .doc, .pdf

    Args:
        filepath: 文件路径

    Returns:
        解析后的文本内容

    Raises:
        UnsupportedFormatError: 不支持的文件格式
        EmptyFileError: 文件内容为空
        MinerUNotInstalledError: 需要 MinerU 但未安装
        MinerUTokenRequiredError: MinerU 需要 API token
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"文件不存在: {filepath}")

    ext = Path(filepath).suffix.lower()

    # 根据扩展名选择解析器
    if ext in (".md", ".txt"):
        text = _parse_txt_md(filepath)
    elif ext == ".docx":
        text = _parse_docx(filepath)
    elif ext == ".doc":
        text = _parse_doc(filepath)
    elif ext == ".pdf":
        # 先检测是否为扫描件
        if is_scanned_pdf(filepath):
            text = _parse_scanned_pdf(filepath)
        else:
            text = _parse_pdf(filepath)
    else:
        raise UnsupportedFormatError()

    if not text or not text.strip():
        raise EmptyFileError(f"文件 {filepath} 解析后内容为空")

    return text


def _parse_txt_md(filepath: str) -> str:
    """解析 .md / .txt 文件"""
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return text


def _parse_docx(filepath: str) -> str:
    """解析 .docx 文件

    先用 python-docx 提取正文段落；如果内容为空（可能是文本框/形状/复杂排版），
    则 fallback 到 MinerU flash-extract 提取文本。
    """
    from docx import Document

    doc = Document(filepath)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    # 如果正文为空，尝试用 MinerU 提取文本框/形状/流程图里的文字
    if not text.strip():
        text = _parse_docx_with_mineru(filepath)

    return text


def _parse_docx_with_mineru(filepath: str) -> str:
    """用 MinerU 解析 .docx（处理文本框、形状、流程图等）"""
    mineru_path = shutil.which("mineru-open-api")
    if not mineru_path:
        return ""

    tmp_dir = tempfile.mkdtemp()
    try:
        result = subprocess.run(
            ["mineru-open-api", "flash-extract", filepath, "-o", tmp_dir],
            capture_output=True,
            text=True,
            timeout=300,
        )
        if result.returncode != 0:
            return ""

        # 查找输出的 .md 文件
        for root, dirs, files in os.walk(tmp_dir):
            for f in files:
                if f.endswith(".md"):
                    md_path = os.path.join(root, f)
                    with open(md_path, "r", encoding="utf-8", errors="ignore") as fh:
                        return fh.read()
        return ""
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def _parse_doc(filepath: str) -> str:
    """解析 .doc 文件

    优先使用 macOS textutil，不可用时尝试 python-docx 作为后备。
    注意：.doc 格式支持有限，建议转换为 .docx。
    """
    # 方案 1: macOS textutil
    if shutil.which("textutil"):
        tmp_dir = tempfile.mkdtemp()
        try:
            tmp_txt = os.path.join(tmp_dir, "output.txt")
            result = subprocess.run(
                ["textutil", "-convert", "txt", "-output", tmp_txt, filepath],
                capture_output=True,
                timeout=30,
            )
            if result.returncode == 0 and os.path.exists(tmp_txt) and os.path.getsize(tmp_txt) > 10:
                with open(tmp_txt, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()
        except Exception:
            pass
        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)

    # 方案 2: python-docx 后备（.doc 兼容性有限）
    try:
        from docx import Document
        doc = Document(filepath)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception:
        return ""


def is_scanned_pdf(filepath: str) -> bool:
    """检测 PDF 是否为扫描件（每页平均文字 < 50 字符）"""
    import fitz

    doc = fitz.open(filepath)
    total_chars = sum(len(page.get_text().strip()) for page in doc)
    total_pages = len(doc)
    doc.close()
    return total_pages > 0 and (total_chars / total_pages) < 50


def _parse_pdf(filepath: str) -> str:
    """解析普通 PDF 文件"""
    import fitz

    doc = fitz.open(filepath)
    texts = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            texts.append(text)
    doc.close()
    return "\n".join(texts)


def _parse_scanned_pdf(filepath: str) -> str:
    """用 MinerU 解析扫描件 PDF"""
    # 1. 检测 MinerU 是否安装
    mineru_path = shutil.which("mineru-open-api")
    if not mineru_path:
        raise MinerUNotInstalledError(
            "需要 MinerU 来解析扫描件 PDF，但未检测到 mineru-open-api",
            hint="安装命令: pip install mineru-open-api\n"
            "MinerU 免费免登录即可使用 flash-extract\n"
            "详细说明: https://github.com/opendatalab/MinerU",
        )

    # 2. 尝试 flash-extract（免 token）
    tmp_dir = tempfile.mkdtemp()
    try:
        result = subprocess.run(
            ["mineru-open-api", "flash-extract", filepath, "-o", tmp_dir],
            capture_output=True,
            text=True,
            timeout=300,
        )

        if result.returncode != 0:
            stderr_hint = result.stderr.strip() if result.stderr else ""
            raise MinerUTokenRequiredError(
                f"MinerU flash-extract 失败: {stderr_hint or '可能文件过大或格式不兼容'}",
                hint="可尝试:\n"
                "1. 使用 mineru-open-api extract --token YOUR_TOKEN（需注册 https://mineru.net）\n"
                "2. 将 PDF 转为图片后先用其他 OCR 工具处理",
            )

        # 3. 查找输出的 .md 文件
        for root, dirs, files in os.walk(tmp_dir):
            for f in files:
                if f.endswith(".md"):
                    md_path = os.path.join(root, f)
                    with open(md_path, "r", encoding="utf-8", errors="ignore") as fh:
                        return fh.read()

        return ""  # 没找到输出文件
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)