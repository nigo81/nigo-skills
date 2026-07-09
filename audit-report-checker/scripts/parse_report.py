#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
parse_report.py — 审计报告解析层（四路分流，只输出原始数据）

设计原则 D1：脚本 MUST NOT 做报表分类识别、续表合并、列映射、
页眉页脚拆分、关键词匹配报表名、同义词归一化。这些是 Claude 主控
的语义工作。脚本只输出原始 markdown + 原始二维表格。

四路分流：
  - PDF 文本型（首页文本 > 50 字符）  → pdfplumber
  - PDF 扫描型                        → mineru-open-api extract --model vlm
  - Word 文本表格（inline_shapes 少）  → python-docx
  - Word 图片报表（inline_shapes 多）  → mineru-open-api extract --model vlm

输出（到 -o 指定目录）：
  - report.md             全文 markdown，每段用 HTML 注释标注来源
  - extracted_tables.json 结构化二维表格数据

用法：
  python3 parse_report.py <文件或目录> -o <输出目录>
  python3 parse_report.py --help
"""
import argparse
import json
import re
import shutil
import subprocess
import sys
from collections import defaultdict
from pathlib import Path

try:
    import openpyxl
except ImportError:
    openpyxl = None

# ──────────────────────────────────────────────────────────────────
# 数字清洗（单元格级）
# ──────────────────────────────────────────────────────────────────
# pdfplumber 常见空格污染：`1 ,767,534.81` / `- 331,089.79` / `2 19,072.58`
# 规则：移除空格类字符后，若结果仅含 [\d,.\-()%] 且至少含一个数字，
#       则返回去空格值；否则保留原文（含中文内容不清洗）。
_NUM_CHARS_RE = re.compile(r"[\d,.\-()%]+")
_HAS_DIGIT_RE = re.compile(r"\d")
_WS_CLEAN_RE = re.compile(r"[ \t\u3000]")  # 普通空格 / tab / 全角空格


def clean_number_cell(text):
    """清洗单元格内数字的空格污染。None → ""。"""
    if text is None:
        return ""
    s = str(text)
    stripped = _WS_CLEAN_RE.sub("", s)
    if stripped and _NUM_CHARS_RE.fullmatch(stripped) and _HAS_DIGIT_RE.search(stripped):
        return stripped
    return s


# ──────────────────────────────────────────────────────────────────
# Excel 报表文件识别（分体式报告：Word/PDF 正文 + Excel 四表）
# ──────────────────────────────────────────────────────────────────
_STATEMENT_SHEET_KEYWORDS = ("资产负债表", "利润表", "现金流量表", "权益变动表", "股东权益")


# Excel 报表文件后缀（含宏启用工作簿 .xlsm）
_XLSX_EXTS = (".xlsx", ".xlsm", ".xlsb", ".XLSX", ".XLSM", ".XLSB")


def _is_statement_xlsx(xlsx_path: Path) -> bool:
    """通过 sheet 名判断 xlsx/xlsm 是否为四表文件。"""
    if openpyxl is None:
        return False
    try:
        wb = openpyxl.load_workbook(str(xlsx_path), read_only=True, data_only=True)
        sheets_text = " ".join(wb.sheetnames)
        return any(k in sheets_text for k in _STATEMENT_SHEET_KEYWORDS)
    except Exception:
        return False


def _collect_xlsx_files(input_path: Path):
    """收集输入目录/文件中的 Excel 文件，返回 [(源路径, 角色), ...]。"""
    if input_path.is_file() and input_path.suffix.lower() in _XLSX_EXTS:
        role = "statements" if _is_statement_xlsx(input_path) else "notes"
        return [(input_path, role)]
    if not input_path.is_dir():
        return []
    files = sorted(
        (f for f in input_path.iterdir() if f.is_file() and f.suffix.lower() in _XLSX_EXTS),
        key=lambda x: x.name,
    )
    result = []
    for f in files:
        role = "statements" if _is_statement_xlsx(f) else "notes"
        result.append((f, role))
    return result

def _escape_md_cell(s):
    """转义 markdown 表格单元格内的特殊字符。"""
    if s is None:
        return ""
    return (str(s)
            .replace("\\", "\\\\")
            .replace("|", "\\|")
            .replace("\n", " ")
            .replace("\r", " ")
            .strip())


def render_md_table(headers, rows):
    """把二维表格渲染成 markdown 表格语法。返回字符串（不含前后空行）。"""
    if not headers and not rows:
        return ""
    ncols = max(len(headers) if headers else 0,
                max((len(r) for r in rows), default=0))
    if ncols == 0:
        return ""
    headers = list(headers) + [""] * (ncols - len(headers))
    norm_rows = [list(r) + [""] * (ncols - len(r)) for r in rows]
    lines = []
    lines.append("| " + " | ".join(_escape_md_cell(h) for h in headers) + " |")
    lines.append("| " + " | ".join("---" for _ in range(ncols)) + " |")
    for r in norm_rows:
        lines.append("| " + " | ".join(_escape_md_cell(c) for c in r) + " |")
    return "\n".join(lines)


def parse_md_tables(md_text):
    """从 markdown 文本解析出表格，返回 [(headers, rows), ...]。

    用于 mineru 输出的 md。第一行作 headers，其余作 rows；跳过分隔行。
    """
    tables = []
    headers = None
    rows = []
    in_table = False
    for line in md_text.splitlines():
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped[1:-1].split("|")]
            # 分隔行 |---|---|
            non_empty = [c for c in cells if c]
            if non_empty and all(re.fullmatch(r"[-:\s]+", c) for c in non_empty):
                in_table = True
                continue
            if not in_table or headers is None:
                headers = cells
                rows = []
                in_table = True
            else:
                rows.append(cells)
        else:
            if in_table and headers is not None:
                tables.append((headers, rows))
                headers = None
                rows = []
                in_table = False
    if in_table and headers is not None:
        tables.append((headers, rows))

    # HTML <table> 解析（mineru vlm 输出 HTML 表格，非管道格式）
    for table_html in re.finditer(r"<table[^>]*>(.*?)</table>", md_text, re.DOTALL):
        trs = []
        for tr in re.finditer(r"<tr[^>]*>(.*?)</tr>", table_html.group(1), re.DOTALL):
            cells = []
            for cell in re.finditer(r"<t[dh][^>]*>(.*?)</t[dh]>", tr.group(1), re.DOTALL):
                cell_text = re.sub(r"<[^>]+>", "", cell.group(1)).strip()
                cells.append(cell_text)
            if cells:
                trs.append(cells)
        if trs:
            tables.append((trs[0], trs[1:] if len(trs) > 1 else []))

    return tables


# ──────────────────────────────────────────────────────────────────
# 元数据提取（编制单位 / 期间，简单正则，非报表分类）
# ──────────────────────────────────────────────────────────────────
_COMPANY_RE = re.compile(r"编制单位\s*[:：]\s*([^\s,，。；;]+)")
_PERIOD_RE = re.compile(r"(\d{4}\s*年度|\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日)")


def _extract_company(text):
    if not text:
        return ""
    m = _COMPANY_RE.search(text)
    return m.group(1).strip() if m else ""


def _extract_period(text):
    if not text:
        return ""
    m = _PERIOD_RE.search(text)
    if not m:
        return ""
    return re.sub(r"\s+", "", m.group(1))


# ──────────────────────────────────────────────────────────────────
# PDF: pdfplumber（文本型）
# ──────────────────────────────────────────────────────────────────
def _text_excluding_tables(page, table_bboxes):
    """提取页面文本，跳过落在表格 bbox 内的字符（避免与表格重复呈现）。

    不做页眉/页脚拆分（D1），只做表格区域去重。
    """
    if not table_bboxes:
        return page.extract_text() or ""
    chars = []
    for ch in page.chars:
        cx = (ch["x0"] + ch["x1"]) / 2
        cy = (ch["top"] + ch["bottom"]) / 2
        in_table = False
        for bb in table_bboxes:
            if bb[0] <= cx <= bb[2] and bb[1] <= cy <= bb[3]:
                in_table = True
                break
        if not in_table:
            chars.append(ch)
    if not chars:
        return ""
    # 按 top 聚合成行（容差 3pt）
    lines = defaultdict(list)
    for ch in chars:
        lines[round(ch["top"] / 3)].append(ch)
    result = []
    for key in sorted(lines.keys()):
        line_chars = sorted(lines[key], key=lambda c: c["x0"])
        result.append("".join(c["text"] for c in line_chars))
    return "\n".join(result)


def is_scanned_pdf(pdf_path):
    """首页文本 <= 50 字符视为扫描型。"""
    import pdfplumber
    try:
        with pdfplumber.open(pdf_path) as pdf:
            if not pdf.pages:
                return False
            txt = pdf.pages[0].extract_text() or ""
            return len(txt.strip()) <= 50
    except Exception:
        return False


def _merge_broken_cn(text):
    """合并 pdfplumber 提取时因字符坐标间距产生的中文断字（'金融机 构'→'金融机构'）。

    只合并 汉字+空白+汉字 模式（循环至稳定）；数字/英文间空格不动，避免破坏正常分隔。
    审计报告正文中文连续无空格，'汉字+空格+汉字'大概率是 pdfplumber 拼接误差（断字），
    合并可减少 DeepSeek 文本检查把断字当错别字的误报。
    """
    if not text:
        return text
    prev = None
    while prev != text:
        prev = text
        text = re.sub(r'([\u4e00-\u9fa5])\s+([\u4e00-\u9fa5])', r'\1\2', text)
    return text


def _process_pdf_text(file_path, rel_path):
    """pdfplumber 提取文本型 PDF。支持混合型：检测连续空白页自动 mineru OCR 补充。"""
    import pdfplumber
    import tempfile

    md_parts = []
    tables_out = []
    company = ""
    period = ""

    # 记录每页的文本量（用于空白页检测）
    page_text_lengths = []  # [(page_no, text_length), ...]
    # 记录每页在 md_parts 中的 SOURCE 注释位置（后续可能替换/补充）
    page_md_index = {}  # page_no -> md_parts index

    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            page_no = i + 1
            found = page.find_tables()  # pdfplumber 0.11 返回 list[Table]
            table_bboxes = [t.bbox for t in found]
            tables_data = [t.extract() for t in found]
            text = _text_excluding_tables(page, table_bboxes)
            text = _merge_broken_cn(text)  # 合并中文断字，减少 DeepSeek 文本检查误报

            # 记录该页的文本量（去除空白后的字符数）
            text_length = len(text.strip()) if text else 0
            page_text_lengths.append((page_no, text_length))

            if not company:
                company = _extract_company(page.extract_text() or text)
            if not period:
                period = _extract_period(page.extract_text() or text)

            # 记录 SOURCE 注释位置
            page_md_index[page_no] = len(md_parts)
            md_parts.append(
                f'<!-- SOURCE file="{rel_path}" page={page_no} method=pdfplumber -->')
            if text.strip():
                md_parts.append(text.strip())

            for tbl in tables_data:
                if not tbl:
                    continue
                headers = [clean_number_cell(c) for c in tbl[0]]
                rows = [[clean_number_cell(c) for c in r] for r in tbl[1:]]
                # 跳过完全空的表格
                if not any(headers) and not any(any(r) for r in rows):
                    continue
                tables_out.append({
                    "id": 0,  # 全局重新编号在 main 里做
                    "name": "",
                    "source_file": rel_path,
                    "page": page_no,
                    "headers": headers,
                    "rows": rows,
                    "method": "pdfplumber",
                })
                rendered = render_md_table(headers, rows)
                if rendered:
                    md_parts.append(rendered)

    # 检测连续空白页（文本量 < 10 字符，且连续 ≥ 3 页）
    blank_ranges = []
    i = 0
    while i < len(page_text_lengths):
        page_no, text_length = page_text_lengths[i]
        if text_length < 10:
            # 开始连续空白页
            start_page = page_no
            end_page = page_no
            # 向后查找连续空白页
            j = i + 1
            while j < len(page_text_lengths):
                next_page_no, next_text_length = page_text_lengths[j]
                if next_text_length < 10 and next_page_no == end_page + 1:
                    end_page = next_page_no
                    j += 1
                else:
                    break
            # 只有连续 ≥ 3 页才记录
            if end_page - start_page + 1 >= 3:
                blank_ranges.append((start_page, end_page))
            i = j
        else:
            i += 1

    # 对连续空白页自动 mineru OCR 补充
    for start_page, end_page in blank_ranges:
        print(f"[警告] page {start_page}-{end_page} 连续{end_page - start_page + 1}页文本为空，"
              f"尝试 mineru OCR 补充...", file=sys.stderr)

        # 调用 mineru OCR（使用临时目录）
        with tempfile.TemporaryDirectory() as tmpdir:
            mineru_result = _mineru_ocr_pages(
                file_path, start_page, end_page, Path(tmpdir))

            if mineru_result is None:
                # mineru 失败，降级为警告
                print(f"[警告] mineru OCR 失败，保留空白页警告", file=sys.stderr)
                # 插入原有警告
                insert_index = page_md_index.get(start_page, -1)
                if insert_index >= 0:
                    warning_comment = (
                        f'<!-- WARNING: page {start_page}-{end_page} 文本为空（连续{end_page - start_page + 1}页），'
                        f'可能为图片/扫描页，pdfplumber 无法提取。'
                        f'若为四表区域，建议用 mineru-open-api extract "{file_path}" '
                        f'--pages {start_page}-{end_page} -o <目录> --model vlm --language ch 重新OCR提取 -->'
                    )
                    visible_warning = (
                        f'⚠️ **第{start_page}-{end_page}页为空白（可能图片/扫描），'
                        f'pdfplumber未能提取。'
                        f'若此处应是财务报表，需用 mineru OCR：'
                        f'`mineru-open-api extract "{file_path}" --pages {start_page}-{end_page} --model vlm`**'
                    )
                    md_parts.insert(insert_index, warning_comment)
                    md_parts.insert(insert_index + 1, visible_warning)
                continue

            # mineru 成功，解析结果
            md_content, tables_list, page_offset = mineru_result

            # 合并 mineru 表格到 tables_out（粗略页码映射：从 start_page 开始递增）
            for idx, tbl in enumerate(tables_list):
                # 粗略估算页码：start_page + idx（mineru 输出可能不精确）
                # 真正的页码由 Claude 主控根据内容复核
                estimated_page = start_page + idx
                tables_out.append({
                    "id": 0,  # 全局重新编号在 main 里做
                    "name": "",
                    "source_file": rel_path,
                    "page": estimated_page,
                    "headers": tbl["headers"],
                    "rows": tbl["rows"],
                    "method": "mineru-vlm",
                })

            # 补充 report.md 内容（替换原有 SOURCE 注释为 mineru 内容）
            # 找到 start_page 在 md_parts 中的位置
            insert_index = page_md_index.get(start_page, -1)
            if insert_index >= 0:
                # 保留 SOURCE 注释（修改 method）
                md_parts[insert_index] = (
                    f'<!-- SOURCE file="{rel_path}" page={start_page} method=mineru-vlm -->'
                )
                # 插入 mineru 内容（在 SOURCE 之后）
                md_parts.insert(insert_index + 1,
                    "> **注意**：此部分由 mineru VLM OCR 提取（自动补充空白页），数字建议人工复核关键勾稽关系。")
                md_parts.insert(insert_index + 2, md_content.strip())

    return {
        "md_parts": md_parts,
        "tables": tables_out,
        "method": "pdfplumber",  # 主方法仍是 pdfplumber，mineru 作为补充
        "company": company,
        "period": period,
    }


# ──────────────────────────────────────────────────────────────────
# Word: python-docx（文本表格）
# ──────────────────────────────────────────────────────────────────
def _docx_shape_table_counts(docx_path):
    """返回 (inline_shapes 数, tables 数)。"""
    from docx import Document
    doc = Document(str(docx_path))
    return len(doc.inline_shapes), len(doc.tables)


def is_image_report_docx(docx_path):
    """判断 Word 是否为图片报表（应走 mineru）。

    判据：inline_shapes >= 3（图片多，疑似报表以图片插入），
    或 表格数 == 0 且有图片。
    """
    try:
        n_shapes, n_tables = _docx_shape_table_counts(docx_path)
    except Exception:
        return False
    return n_shapes >= 3 or (n_tables == 0 and n_shapes > 0)


def _sdt_aware_extract_tbl(tbl_elem):
    """sdt-aware 提取 Word 表格行（处理 sdt 内容控件包裹的 tc + gridSpan 横向合并）。

    解决两类 python-docx 漏读：
    - sdt 包裹：项目名等 tc 在 <w:sdt><w:sdtContent><w:tc> 内，python-docx
      的 row.cells 用 tr.findall(w:tc) 不递归进 sdtContent，会漏读被 sdt
      包裹的单元格（典型表现：资产负债表项目名列全空、金额错位）。
    - gridSpan：横向合并单元格按跨度 N 重复填充（[txt]*gs），保持行内列数
      与表头对齐，行为与原 python-docx row.cells 一致。

    注意：调用方应确保 tbl_elem 是叶子表（不含嵌套 w:tbl），否则 tr.iter(w:tc)
    会误抓嵌套表的单元格。嵌套表场景由调用方先筛选叶子后代。
    """
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    rows = []
    for tr in tbl_elem.iter(f"{W}tr"):
        cells = []
        for tc in tr.iter(f"{W}tc"):
            gs = 1
            tcPr = tc.find(f"{W}tcPr")
            if tcPr is not None:
                gs_el = tcPr.find(f"{W}gridSpan")
                if gs_el is not None and gs_el.get(f"{W}val"):
                    try:
                        gs = int(gs_el.get(f"{W}val"))
                    except (ValueError, TypeError):
                        gs = 1
            txt = "".join(t.text or "" for t in tc.iter(f"{W}t"))
            txt = clean_number_cell(txt)
            cells.extend([txt] * gs)
        rows.append(cells)
    return rows


def _process_docx_text(file_path, rel_path):
    """python-docx 提取 Word 文本表格。按文档原始顺序遍历段落与表格。"""
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(str(file_path))
    md_parts = [f'<!-- SOURCE file="{rel_path}" method=python-docx -->']
    tables_out = []
    company = ""
    period = ""
    # 章节跟踪（Word 无页码，用章节定位）：当前大章节 + 附注五明细科目序号
    _MAJOR_RE = re.compile(r'^([一二三四五六七八九十]+)[、．.]\s*(\S.{0,30})')
    _NOTE5_ITEM_RE = re.compile(r'^(\d{1,2})[．、.]\s*([\u4e00-\u9fa5][\u4e00-\u9fa5（）()A-Za-z0-9]{1,20})')
    _CN_NUM = "一二三四五六七八九十"  # 大章节中文数字映射（按 Heading1 出现顺序）
    current_major = ""         # 当前大章节中文数字（如"五"）
    current_major_title = ""   # 大章节标题
    major_seq = 0              # Heading1 大章节计数器（映射到中文数字）
    note5_seq = 0              # 附注五明细科目序号
    note5_subject = ""         # 当前附注五明细科目名
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag
        if tag.endswith("}p"):
            para = Paragraph(child, doc)
            text = para.text.strip()
            if not text:
                continue
            if not company:
                company = _extract_company(text)
            if not period:
                period = _extract_period(text)
            # 章节跟踪：大章节优先 Heading1 样式（金星实测：大章节是 Heading1 无"五、"编号），兜底"中文数字、"文字
            style_name = para.style.name if para.style else ""
            is_heading1 = "Heading 1" in style_name or "标题 1" in style_name or "标题1" in style_name
            if is_heading1:
                major_seq += 1
                current_major = _CN_NUM[major_seq - 1] if major_seq <= len(_CN_NUM) else str(major_seq)
                current_major_title = text
                note5_seq = 0  # 进入新大章节，重置明细序号
            else:
                m_major = _MAJOR_RE.match(text)  # 兜底：其他报告可能用"五、"文字编号
                if m_major:
                    cn = m_major.group(1)
                    current_major = cn
                    current_major_title = m_major.group(2).strip()
                    major_seq = _CN_NUM.find(cn) + 1 if cn in _CN_NUM else major_seq
                    note5_seq = 0
                elif current_major == "五":
                    # 明细科目识别（通用：不同事务所样式不同——金星Heading2、博达Normal纯文字）
                    if "Heading 2" in style_name or "标题 2" in style_name or "标题2" in style_name:
                        note5_seq += 1; note5_subject = text
                    elif not text.startswith(("（", "(", "本公司", "注", "年末", "年初", "本年", "上述", "截止")):
                        m_item = _NOTE5_ITEM_RE.match(text)
                        if m_item:
                            note5_seq += 1; note5_subject = m_item.group(2).strip()
                        elif 2 <= len(text) <= 10 and re.search(r'[\u4e00-\u9fa5]', text) and not re.search(r'[，。；：、,;:0-9]', text):
                            # 短文本无标点（通用兜底：博达Normal"货币资金"/"应收票据"等纯文字科目名）
                            note5_seq += 1; note5_subject = text
            md_parts.append(text)
        elif tag.endswith("}tbl") or tag.endswith("}sdt"):
            W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            # 检测后代 tbl：四表可能被包裹在容器表或 sdt 内容控件内。
            # body 直接子级的 w:sdt 是常见包裹方式（巨东报告四表嵌在 sdt 内），
            # python-docx 与 body.iterchildren 都不会进入 sdt，导致四表漏读。
            descendants = child.findall(f".//{W_NS}tbl")
            if descendants:
                # 容器（tbl/sdt）：只提取叶子后代（不含自己的 .//w:tbl），跳过容器
                # 本身（容器 tc 含嵌套表时 iter(w:tc) 会串列）。
                targets = [t for t in descendants if not t.findall(f".//{W_NS}tbl")]
                method_tag = "python-docx-nested"
            elif tag.endswith("}tbl"):
                # 顶层叶子表：提取自身
                targets = [child]
                method_tag = "python-docx"
            else:
                continue  # sdt 无后代 tbl（纯文本/段落 sdt），跳过
            for tbl_elem in targets:
                rows = _sdt_aware_extract_tbl(tbl_elem)
                if not rows:
                    continue
                headers = rows[0]
                data_rows = rows[1:]
                if not any(headers) and not any(any(r) for r in data_rows):
                    continue
                # chapter 归属（Word 章节定位）：附注五下 → "附注五-N 科目"；其他 → 大章节标题。
                # 嵌套四表继承遍历到容器表时的章节上下文（财务报表章节）。
                if current_major == "五" and note5_seq > 0:
                    chapter = f"附注五-{note5_seq} {note5_subject}".strip()
                elif current_major:
                    chapter = current_major_title
                else:
                    chapter = ""
                tables_out.append({
                    "id": 0,
                    "name": note5_subject if current_major == "五" else "",
                    "source_file": rel_path,
                    "page": None,
                    "chapter": chapter,
                    "headers": headers,
                    "rows": data_rows,
                    "method": method_tag,
                })
                rendered = render_md_table(headers, data_rows)
                if rendered:
                    md_parts.append(rendered)
    # 四表齐全性自检：关键项目缺失时输出 stderr 诊断（不静默跳过）
    _FOUR_STMT_KEYS = ["资产总计", "流动资产合计", "负债和所有者权益", "营业利润"]
    _all_cell_text = " ".join(
        str(c) for t in tables_out
        for r in ([t["headers"]] + t.get("rows", [])) for c in r
    )
    for _key in _FOUR_STMT_KEYS:
        if _key not in _all_cell_text:
            print(f"[WARN] 四表关键项目「{_key}」缺失，疑似表格被嵌套/sdt 包裹未提取", file=sys.stderr)

    return {
        "md_parts": md_parts,
        "tables": tables_out,
        "method": "python-docx",
        "company": company,
        "period": period,
    }


# ──────────────────────────────────────────────────────────────────
# mineru 云端 API（扫描 PDF / Word 图片报表）
# ──────────────────────────────────────────────────────────────────
def _mineru_ocr_one_segment(pdf_path, seg_start, seg_end, mineru_out, use_fitz=True):
    """对单段页码范围调一次 mineru VLM，返回 md_content 或 None。"""
    if use_fitz:
        try:
            import fitz
            split_pdf = mineru_out.parent / f"_split_p{seg_start}-{seg_end}.pdf"
            doc = fitz.open(str(pdf_path))
            new = fitz.open()
            new.insert_pdf(doc, from_page=seg_start - 1, to_page=seg_end - 1)
            new.save(str(split_pdf))
            new.close()
            doc.close()
            size_mb = split_pdf.stat().st_size / 1024 / 1024
            print(f"[INFO] 拆分 page {seg_start}-{seg_end} 为 {size_mb:.2f}MB",
                  file=sys.stderr)
            target = split_pdf
            pages_arg = None
        except Exception as e:
            print(f"[INFO] pymupdf 拆分失败({e})，用 --pages", file=sys.stderr)
            target = pdf_path
            pages_arg = f"{seg_start}-{seg_end}"
    else:
        target = pdf_path
        pages_arg = f"{seg_start}-{seg_end}"

    seg_out = mineru_out / f"seg_{seg_start}-{seg_end}"
    seg_out.mkdir(parents=True, exist_ok=True)
    cmd = ["mineru-open-api", "extract", str(target), "-o", str(seg_out),
           "-f", "md", "--model", "vlm", "--language", "ch"]
    if pages_arg:
        cmd += ["--pages", pages_arg]
    print(f"[INFO] mineru VLM 提取 page {seg_start}-{seg_end}", file=sys.stderr)

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=1800)
    except subprocess.TimeoutExpired:
        print(f"[ERROR] mineru 超时: page {seg_start}-{seg_end}", file=sys.stderr)
        return None
    except FileNotFoundError:
        print("[ERROR] mineru-open-api 未安装。", file=sys.stderr)
        return None

    if result.returncode != 0:
        print(f"[ERROR] mineru 失败 (rc={result.returncode}): page {seg_start}-{seg_end}",
              file=sys.stderr)
        print(f"  stderr: {result.stderr[:500]}", file=sys.stderr)
        return None

    md_files = list(seg_out.rglob("*.md"))
    if not md_files:
        print(f"[ERROR] mineru 未输出 md: page {seg_start}-{seg_end}", file=sys.stderr)
        return None
    return max(md_files, key=lambda f: f.stat().st_size).read_text(encoding="utf-8")


def _mineru_ocr_pages(pdf_path, page_start, page_end, out_dir):
    """调用 mineru-open-api 对页码范围 OCR。大范围自动分段（每段≤4页，避免OSS上传broken pipe）。

    Returns:
        (md_content, tables_list, page_offset) 或 None
    """
    if shutil.which("mineru-open-api") is None:
        print("[ERROR] mineru-open-api 未安装。请运行: "
              "npm install -g mineru-open-api  &&  mineru-open-api auth", file=sys.stderr)
        return None

    safe_stem = re.sub(r"[^\w\u4e00-\u9fff]", "_", pdf_path.stem)
    mineru_out = out_dir / f"_mineru_{safe_stem}_p{page_start}-{page_end}"
    mineru_out.mkdir(parents=True, exist_ok=True)

    print(f"[隐私提示] page {page_start}-{page_end} 为图片页，将调用 mineru 云端 OCR "
          f"（数据上传 mineru.net），如需跳过请中断", file=sys.stderr)

    # 分段：每段最多 4 页（实测 12 页 4.16MB 会 OSS broken pipe，4 页约 1.4MB 稳定）
    MAX_PAGES_PER_SEGMENT = 4
    segments = []
    cur = page_start
    while cur <= page_end:
        seg_end = min(cur + MAX_PAGES_PER_SEGMENT - 1, page_end)
        segments.append((cur, seg_end))
        cur = seg_end + 1

    if len(segments) > 1:
        print(f"[INFO] page {page_start}-{page_end} 拆为 {len(segments)} 段 OCR "
              f"（每段≤{MAX_PAGES_PER_SEGMENT}页，避免上传超时）", file=sys.stderr)

    all_md = []
    any_success = False
    for seg_start, seg_end in segments:
        md = _mineru_ocr_one_segment(pdf_path, seg_start, seg_end, mineru_out)
        if md:
            all_md.append(md)
            any_success = True
        else:
            print(f"[警告] page {seg_start}-{seg_end} OCR 失败，跳过该段", file=sys.stderr)

    if not any_success:
        return None

    md_content = "\n\n".join(all_md)

    # 解析表格
    tables_list = []
    for headers, rows in parse_md_tables(md_content):
        headers = [clean_number_cell(h) for h in headers]
        rows = [[clean_number_cell(c) for c in r] for r in rows]
        if not any(headers) and not any(any(r) for r in rows):
            continue
        tables_list.append({"headers": headers, "rows": rows})

    return (md_content, tables_list, page_start)


def _process_mineru(file_path, rel_path, out_dir):
    """调用 mineru-open-api extract --model vlm 提取，返回 markdown。"""
    if shutil.which("mineru-open-api") is None:
        print("[ERROR] mineru-open-api 未安装。请运行: "
              "npm install -g mineru-open-api  &&  mineru-open-api auth",
              file=sys.stderr)
        return None

    # 每个文件独立子目录，避免输出冲突
    safe_stem = re.sub(r"[^\w\u4e00-\u9fff]", "_", file_path.stem)
    mineru_out = out_dir / f"_mineru_{safe_stem}"
    mineru_out.mkdir(parents=True, exist_ok=True)

    cmd = [
        "mineru-open-api", "extract",
        str(file_path),
        "-o", str(mineru_out),
        "-f", "md",
        "--model", "vlm",
        "--language", "ch",
    ]
    print(f"[INFO] 调用 mineru VLM 提取: {rel_path}", file=sys.stderr)
    try:
        result = subprocess.run(
            cmd, capture_output=True, text=True, timeout=1800)
    except subprocess.TimeoutExpired:
        print(f"[ERROR] mineru 超时（30 分钟）: {rel_path}", file=sys.stderr)
        return None
    except FileNotFoundError:
        print("[ERROR] mineru-open-api 未安装。请运行: "
              "npm install -g mineru-open-api  &&  mineru-open-api auth",
              file=sys.stderr)
        return None

    if result.returncode != 0:
        print(f"[ERROR] mineru 失败 (rc={result.returncode}): {rel_path}",
              file=sys.stderr)
        print(f"  stdout: {result.stdout[:800]}", file=sys.stderr)
        print(f"  stderr: {result.stderr[:800]}", file=sys.stderr)
        return None

    md_files = list(mineru_out.rglob("*.md"))
    if not md_files:
        print(f"[ERROR] mineru 未输出 md 文件: {rel_path}", file=sys.stderr)
        print(f"  stdout: {result.stdout[:800]}", file=sys.stderr)
        return None
    md_file = max(md_files, key=lambda f: f.stat().st_size)
    md_content = md_file.read_text(encoding="utf-8")

    # 从 mineru markdown 解析表格
    tables_out = []
    for headers, rows in parse_md_tables(md_content):
        headers = [clean_number_cell(h) for h in headers]
        rows = [[clean_number_cell(c) for c in r] for r in rows]
        if not any(headers) and not any(any(r) for r in rows):
            continue
        tables_out.append({
            "id": 0,
            "name": "",
            "source_file": rel_path,
            "page": None,  # mineru md 不带稳定页码，留空
            "headers": headers,
            "rows": rows,
            "method": "mineru-vlm",
        })

    company = _extract_company(md_content)
    period = _extract_period(md_content)

    md_parts = [
        f'<!-- SOURCE file="{rel_path}" method=mineru-vlm -->',
        "> **注意**：此部分由 mineru VLM OCR 提取，数字建议人工复核关键勾稽关系。",
        md_content.strip(),
    ]
    return {
        "md_parts": md_parts,
        "tables": tables_out,
        "method": "mineru-vlm",
        "company": company,
        "period": period,
    }


# ──────────────────────────────────────────────────────────────────
# 主流程
# ──────────────────────────────────────────────────────────────────
_SUPPORTED_EXT = (".pdf", ".docx", ".doc")


def _process_file(file_path, rel_path, out_dir):
    """按文件类型分流处理单个文件。返回 result dict 或 None。"""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        if is_scanned_pdf(file_path):
            return _process_mineru(file_path, rel_path, out_dir)
        return _process_pdf_text(file_path, rel_path)
    if suffix == ".docx":
        if is_image_report_docx(file_path):
            return _process_mineru(file_path, rel_path, out_dir)
        return _process_docx_text(file_path, rel_path)
    if suffix == ".doc":
        print(f"[WARN] 跳过 .doc（请另存为 .docx）: {rel_path}", file=sys.stderr)
        return None
    print(f"[WARN] 不支持的格式，跳过: {rel_path}", file=sys.stderr)
    return None


def _collect_files(input_path):
    """返回 (files, root)。目录输入按文件名排序遍历。

    通用去重：同 basename 的 .pdf/.docx/.doc 只解析一份（优先 .docx——
    Word 内容准无断字+表格精确，其次 .pdf 有 page 页码，.doc 不优先处理时跳过）。
    避免同一报告 docx+pdf 双份解析产生重复表 + docx 表 page=null。
    """
    if input_path.is_dir():
        all_files = sorted(
            (f for f in input_path.iterdir()
             if f.is_file() and f.suffix.lower() in _SUPPORTED_EXT),
            key=lambda x: x.name,
        )
        # 同 basename 去重：优先 .docx（Word 内容准无断字+表格精确），其次 .pdf（有页码），.doc 不优先(处理时跳过)
        _PRI = {'.docx': 3, '.pdf': 2, '.doc': 1}
        chosen = {}
        for f in all_files:
            base = f.stem.lower()
            if base not in chosen:
                chosen[base] = f
            elif _PRI.get(f.suffix.lower(), 0) > _PRI.get(chosen[base].suffix.lower(), 0):
                chosen[base] = f  # 优先级高的替换
        files_set = set(chosen.values())
        files = [f for f in all_files if f in files_set]
        for f in all_files:
            if f not in files_set:
                print(f"[INFO] 去重跳过（同 basename 优先 .docx）: {f.name}", file=sys.stderr)
        return files, input_path
    return [input_path], input_path.parent


def main():
    parser = argparse.ArgumentParser(
        description="审计报告解析层：四路分流提取原始 markdown + 原始二维表格（不做分类/合并/列映射）")
    parser.add_argument("input", help="审计报告文件或目录（目录用于分体式报告）")
    parser.add_argument("-o", "--output", required=True, help="输出目录")
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    if not input_path.exists():
        print(f"[ERROR] 输入不存在: {input_path}", file=sys.stderr)
        sys.exit(1)

    out_dir = Path(args.output).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    files, root = _collect_files(input_path)

    # 依赖预检：按待处理文件类型提示安装对应库
    needed = {}
    for f in files:
        ext = f.suffix.lower()
        if ext == ".pdf":
            needed["pdfplumber"] = "pdfplumber"
        elif ext == ".docx":
            needed["python-docx"] = "docx"
    xlsx_files = _collect_xlsx_files(input_path)
    if xlsx_files:
        needed["openpyxl"] = "openpyxl"
    missing = []
    for pkg, mod in needed.items():
        try:
            __import__(mod)
        except ImportError:
            missing.append(pkg)
    if missing:
        print(
            f"[ERROR] 缺少必要依赖，请执行: pip install {' '.join(missing)}",
            file=sys.stderr,
        )
        sys.exit(1)

    # 收集并复制 Excel 文件（分体式报告常见：Word 正文/附注 + Excel 四表）
    files_index = {}
    for xf, role in xlsx_files:
        dst = out_dir / xf.name
        shutil.copy2(str(xf), str(dst))
        files_index[xf.name] = {"role": role, "loader": "xlsx"}
        print(f"[INFO] 复制 Excel 文件: {xf.name} (role={role})", file=sys.stderr)

    if not files and not files_index:
        print(f"[ERROR] 未找到可处理的文件（支持 .pdf/.docx/.xlsx/.xlsm）: {input_path}",
              file=sys.stderr)
        sys.exit(1)

    all_md_parts = []
    all_tables = []
    methods = {}
    sources = []
    company = ""
    period = ""

    for f in files:
        rel = (str(f.relative_to(root)) if input_path.is_dir()
               else f.name)
        print(f"[INFO] 处理: {rel}", file=sys.stderr)
        try:
            result = _process_file(f, rel, out_dir)
        except Exception as e:  # 单文件失败不影响其他文件
            print(f"[ERROR] 处理失败 {rel}: {e}", file=sys.stderr)
            import traceback
            traceback.print_exc()
            continue
        if result is None:
            continue
        all_md_parts.extend(result["md_parts"])
        all_tables.extend(result["tables"])
        methods[rel] = result["method"]
        sources.append(rel)
        if not company and result["company"]:
            company = result["company"]
        if not period and result["period"]:
            period = result["period"]

    # 表格全局重新编号
    for i, t in enumerate(all_tables):
        t["id"] = i

    # 写 report.md
    report_md = out_dir / "report.md"
    report_md.write_text("\n\n".join(all_md_parts) + "\n", encoding="utf-8")

    # 写 extracted_tables.json
    json_obj = {
        "source_root": str(root),  # 原报告根目录的绝对路径，用于还原相对路径
        "sources": sources,
        "company": company,
        "period": period,
        "extraction_methods": methods,
        "tables": all_tables,
    }
    json_path = out_dir / "extracted_tables.json"
    json_path.write_text(
        json.dumps(json_obj, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"[OK] {report_md}", file=sys.stderr)
    print(f"[OK] {json_path}", file=sys.stderr)

    # 写 files_index.json（用于 run_check 识别分体式报告中的 Excel 报表文件）
    if files_index:
        index_path = out_dir / "files_index.json"
        index_path.write_text(
            json.dumps(files_index, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"[OK] {index_path}", file=sys.stderr)

    print(f"     来源文件 {len(sources)} 个，表格 {len(all_tables)} 张",
          file=sys.stderr)
    print(f"     公司: {company or '(未识别)'}  期间: {period or '(未识别)'}",
          file=sys.stderr)


if __name__ == "__main__":
    main()
