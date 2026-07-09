#!/usr/bin/env python3
"""国家规章库 (www.gov.cn/zhengce/xxgk/gjgzk) API 客户端"""

import argparse
import json
import re
import sys
import time
import urllib.request
import urllib.error

SEARCH_URL = "https://sousuoht.www.gov.cn/athena/forward/BD8730CDDA12515E2D9E1B21AA11C0D6"

# Auth key cache file
import os
_KEY_CACHE = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".athena_key_cache.json")

def _load_cached_key() -> str:
    """Load cached athenaappkey, return empty string if expired or missing."""
    try:
        import time
        with open(_KEY_CACHE, "r") as f:
            cache = json.load(f)
        # Cache valid for 1 hour
        if time.time() - cache.get("ts", 0) < 3600:
            return cache["key"]
    except Exception:
        pass
    return ""


def _save_cached_key(key: str):
    import time
    with open(_KEY_CACHE, "w") as f:
        json.dump({"key": key, "ts": time.time()}, f)


def _ensure_playwright_cli():
    """Check if playwright-cli is installed, install if missing."""
    import subprocess, shutil
    if shutil.which("playwright-cli"):
        return True
    print("  playwright-cli 未安装，正在自动安装...", file=sys.stderr)
    try:
        result = subprocess.run(
            ["npm", "install", "-g", "@playwright/cli"],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode == 0 and shutil.which("playwright-cli"):
            print("  playwright-cli 安装成功", file=sys.stderr)
            return True
        print(f"  Warning: playwright-cli 安装失败: {result.stderr.strip()}", file=sys.stderr)
    except Exception as e:
        print(f"  Warning: playwright-cli 安装失败: {e}", file=sys.stderr)
    return False


def _fetch_fresh_key() -> str:
    """Use playwright-cli to get a fresh athenaappkey by intercepting browser requests."""
    import subprocess
    if not _ensure_playwright_cli():
        return ""
    try:
        subprocess.run(["playwright-cli", "open", "https://www.gov.cn/zhengce/xxgk/gjgzk/index.htm"],
                       capture_output=True, timeout=30)
        result = subprocess.run(
            ["playwright-cli", "--raw", "run-code",
             '''async page => {
  const keys = [];
  await page.route(url => url.href.includes('athena/forward'), async route => {
    const key = route.request().headers()['athenaappkey'] || '';
    if (key) keys.push(key);
    await route.continue();
  });
  await page.reload({ waitUntil: 'networkidle' });
  await page.waitForTimeout(3000);
  await page.unroute(url => url.href.includes('athena/forward'));
  return keys.length > 0 ? keys[0] : '';
}'''],
            capture_output=True, text=True, timeout=60
        )
        subprocess.run(["playwright-cli", "close"], capture_output=True, timeout=10)
        output = (result.stdout + result.stderr).strip()
        # --raw run-code returns the value as a quoted string, e.g. "key_value"
        # It may also contain multiple comma-separated identical keys
        for line in output.strip().split("\n"):
            line = line.strip().strip('"')
            if line and len(line) > 20 and "athena" not in line.lower():
                key = line.split(",")[0].strip().strip('"')
                if key:
                    return key
    except Exception as e:
        print(f"Warning: Failed to fetch fresh key: {e}", file=sys.stderr)
        try:
            subprocess.run(["playwright-cli", "close"], capture_output=True, timeout=5)
        except Exception:
            pass
    return ""


def _verify_key(key: str) -> bool:
    """Verify that an athenaappkey is valid by making a small test request."""
    try:
        body = json.dumps({
            "code": "18258ab0ac9",
            "searchFields": [{"fieldName": F_TITLE, "searchWord": "测试", "withHighLight": False}],
            "sorts": [],
            "resultFields": [F_TITLE],
            "trackTotalHits": "true",
            "tableName": "t_1860c735d31",
            "pageSize": 1,
            "pageNo": 1,
            "granularity": "ALL",
        }, ensure_ascii=False).encode("utf-8")
        req = urllib.request.Request(SEARCH_URL, data=body, method="POST")
        req.add_header("Content-Type", "application/json;charset=UTF-8")
        req.add_header("Origin", "https://www.gov.cn")
        req.add_header("Referer", "https://www.gov.cn/")
        req.add_header("athenaappname", "%E8%A7%84%E7%AB%A0%E5%BA%93")
        req.add_header("athenaappkey", key)
        with urllib.request.urlopen(req, timeout=10) as resp:
            result = json.loads(resp.read().decode("utf-8"))
        rc = result.get("resultCode", {})
        code = str(rc.get("code", ""))
        if code.startswith("athena_"):
            return False
        return code == "200"
    except Exception:
        return False


def _get_auth_key() -> str:
    """Get a valid athenaappkey, refreshing if needed."""
    key = _load_cached_key()
    if key and _verify_key(key):
        return key
    print("  athenaappkey 已过期，正在通过浏览器刷新...", file=sys.stderr)
    key = _fetch_fresh_key()
    if key and _verify_key(key):
        _save_cached_key(key)
        print("  athenaappkey 已刷新并验证通过", file=sys.stderr)
    elif key:
        print("  Warning: 刷新的 athenaappkey 验证失败，key 无效", file=sys.stderr)
        return ""
    else:
        print("  Warning: 无法刷新 athenaappkey", file=sys.stderr)
    return key


def _get_auth_headers() -> dict:
    key = _get_auth_key()
    return {
        "Content-Type": "application/json;charset=UTF-8",
        "Origin": "https://www.gov.cn",
        "Referer": "https://www.gov.cn/",
        "athenaappname": "%E8%A7%84%E7%AB%A0%E5%BA%93",
        "athenaappkey": key,
    }

# Field name mapping
F_TITLE = "f_202321360426"
F_TYPE = "f_202321807875"
F_SXX = "f_202321864401"
F_DATE = "f_202321915922"
F_SOURCE = "f_202323394765"
F_ORG = "f_202355832506"
F_ORG2 = "f_202328191239"
F_CONTENT = "f_202321758948"
F_DOC_NUM = "f_202344311304"
F_URL = "doc_pub_url"

RESULT_FIELDS = [
    F_TITLE, F_TYPE, F_SXX, F_DATE, F_SOURCE, F_ORG, F_ORG2,
    F_DOC_NUM, F_URL, "f_20232124962", "f_202321124775",
    "f_202321159816", "f_2023425676953", "f_2023425808265",
    "f_202321136868", "f_20232380533", "f_20232151076",
]


def _clean_highlight(text):
    if not text:
        return ""
    return re.sub(r"</?em[^>]*>", "", str(text))


def _get_field(item, field):
    val = item.get(field, "")
    if isinstance(val, list):
        return val[0] if val else ""
    return val or ""


def search(keyword: str, reg_type: str = None, content_search: bool = False,
           page: int = 1, size: int = 10) -> dict:
    """搜索规章
    
    Args:
        content_search: True=按全文内容搜索, False=按标题搜索(默认)
    """
    search_fields = [
        {"fieldName": F_TYPE, "searchType": "TERM", "withHighLight": True},
        {"fieldName": F_TITLE, "withHighLight": True},
        {"fieldName": F_CONTENT, "withHighLight": True},
        {"fieldName": "f_202321423473", "searchType": "TERM", "withHighLight": True},
        {"fieldName": "f_202321159816", "searchWord": "", "searchType": "TERM"},
        {"fieldName": "f_20232380533", "searchType": "TERM", "withHighLight": True},
        {"fieldName": F_ORG2, "withHighLight": True, "searchType": "TERM"},
        {"fieldName": "f_20221110222856", "withHighLight": True, "searchType": "TERM"},
    ]
    # Set keyword on the right field
    if content_search:
        search_fields[2]["searchWord"] = keyword  # F_CONTENT
    else:
        search_fields[1]["searchWord"] = keyword  # F_TITLE

    if reg_type:
        search_fields[0]["searchWord"] = reg_type

    body = {
        "code": "18258ab0ac9",
        "searchFields": search_fields,
        "sorts": [{}, {"sortField": F_DATE, "sortOrder": "DESC"}],
        "resultFields": RESULT_FIELDS,
        "trackTotalHits": "true",
        "tableName": "t_1860c735d31",
        "pageSize": size,
        "pageNo": page,
        "granularity": "ALL",
    }

    data = json.dumps(body, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request(SEARCH_URL, data=data, method="POST")
    headers = _get_auth_headers()
    for k, v in headers.items():
        req.add_header(k, v)

    with urllib.request.urlopen(req, timeout=30) as resp:
        result = json.loads(resp.read().decode("utf-8"))

    # Parse results — handle both dict and list formats for "data"
    result_obj = result.get("result", {})
    data = result_obj.get("data", {})
    if isinstance(data, list):
        pager = {"total": result_obj.get("totalCount", 0)}
        items = data
    else:
        pager = data.get("pager", {})
        items = data.get("list", [])

    parsed = []
    for item in items:
        parsed.append({
            "title": _clean_highlight(_get_field(item, F_TITLE)),
            "type": _clean_highlight(_get_field(item, F_TYPE)),
            "sxx": _get_field(item, F_SXX),
            "date": _get_field(item, F_DATE).split(" ")[0] if _get_field(item, F_DATE) else "",
            "org": _get_field(item, "f_20232151076") or _get_field(item, F_ORG) or _get_field(item, F_ORG2),
            "source": _get_field(item, F_SOURCE),
            "doc_num": _get_field(item, F_DOC_NUM),
            "url": _get_field(item, F_URL),
        })

    return {"total": pager.get("total", 0), "rows": parsed}


def new_since(date_str: str, reg_type: str = None, page_size: int = 20,
              delay: float = 0.5) -> dict:
    """查找指定日期后新发布的规章（自动翻页，按日期过滤）"""
    all_rows = []
    page = 1
    while True:
        resp = search("", reg_type=reg_type, page=page, size=page_size)
        rows = resp.get("rows", [])
        total = resp.get("total", 0)
        if not rows:
            break
        for row in rows:
            if row.get("date", "") >= date_str:
                all_rows.append(row)
            else:
                # Results are sorted by date desc, so we can stop
                return {"total": len(all_rows), "rows": all_rows}
        page += 1
        time.sleep(delay)
        if page > 1:
            print(f"  翻页: {page} (已获取 {len(all_rows)})", file=sys.stderr)
    return {"total": len(all_rows), "rows": all_rows}


def search_all_pages(keyword: str, reg_type: str = None, content_search: bool = False,
                     max_pages: int = 50, page_size: int = 20, delay: float = 0.5) -> list:
    """搜索所有页的结果，自动翻页"""
    all_rows = []
    page = 1
    while page <= max_pages:
        resp = search(keyword, reg_type=reg_type, content_search=content_search,
                      page=page, size=page_size)
        rows = resp.get("rows", [])
        total = resp.get("total", 0)
        all_rows.extend(rows)
        if len(all_rows) >= total or not rows:
            break
        page += 1
        if page <= max_pages and len(all_rows) < total:
            time.sleep(delay)
            print(f"  翻页: {page}/{(total + page_size - 1) // page_size} (已获取 {len(all_rows)}/{total})", file=sys.stderr)
    return {"total": total, "rows": all_rows}


def batch_check(names: list, delay: float = 0.5) -> list:
    """批量检查规章状态"""
    results = []
    total = len(names)
    for i, name in enumerate(names):
        name = name.strip()
        if not name:
            continue
        clean = re.sub(r'[《》<>〈〉]', '', name).strip()
        try:
            resp = search(clean, size=5)
            rows = resp.get("rows", [])
            best = None
            for row in rows:
                title = row.get("title", "")
                title_clean = re.sub(r'[《》<>〈〉]', '', title).strip()
                if title_clean == clean:
                    # Exact match - prefer 部门规章 over 地方政府规章
                    if best is None or row.get("type") == "部门规章":
                        best = row
            if best is None:
                # Fallback: containment match with length check
                for row in rows:
                    title = row.get("title", "")
                    title_clean = re.sub(r'[《》<>〈〉]', '', title).strip()
                    if clean in title_clean and len(clean) > len(title_clean) * 0.5:
                        if best is None or row.get("type") == "部门规章":
                            best = row
            if best:
                results.append({
                    "name": name,
                    "found": True,
                    "title": best["title"],
                    "type": best.get("type", ""),
                    "status": "有效",  # 国家规章库只收录现行有效规章
                    "date": best.get("date", ""),
                    "org": best.get("org", ""),
                    "url": best.get("url", ""),
                })
            else:
                results.append({"name": name, "found": False, "status": "未找到"})
        except Exception as e:
            results.append({"name": name, "found": False, "status": f"查询失败: {e}"})
        if i < total - 1:
            time.sleep(delay)
        if (i + 1) % 10 == 0 or i == total - 1:
            print(f"  进度: {i+1}/{total}", file=sys.stderr)
    return results


def search_with_content(keyword: str, reg_type: str = None, size: int = 5) -> dict:
    """搜索规章并返回全文内容（用于下载）"""
    search_fields = [
        {"fieldName": F_TYPE, "searchType": "TERM", "withHighLight": True},
        {"fieldName": F_TITLE, "searchWord": keyword, "withHighLight": False},
        {"fieldName": F_CONTENT, "withHighLight": False},
        {"fieldName": "f_202321423473", "searchType": "TERM"},
        {"fieldName": "f_202321159816", "searchWord": "", "searchType": "TERM"},
        {"fieldName": "f_20232380533", "searchType": "TERM"},
        {"fieldName": F_ORG2, "searchType": "TERM"},
    ]
    if reg_type:
        search_fields[0]["searchWord"] = reg_type

    body = {
        "code": "18258ab0ac9",
        "searchFields": search_fields,
        "sorts": [{}, {"sortField": F_DATE, "sortOrder": "DESC"}],
        "resultFields": RESULT_FIELDS + [F_CONTENT],
        "trackTotalHits": "true",
        "tableName": "t_1860c735d31",
        "pageSize": size,
        "pageNo": 1,
        "granularity": "ALL",
    }

    data = json.dumps(body, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request(SEARCH_URL, data=data, method="POST")
    headers = _get_auth_headers()
    for k, v in headers.items():
        req.add_header(k, v)

    with urllib.request.urlopen(req, timeout=30) as resp:
        result = json.loads(resp.read().decode("utf-8"))

    result_obj = result.get("result", {})
    data = result_obj.get("data", {})
    if isinstance(data, list):
        items = data
        total = result_obj.get("totalCount", 0)
    else:
        items = data.get("list", [])
        total = data.get("pager", {}).get("total", 0)

    parsed = []
    for item in items:
        parsed.append({
            "title": _clean_highlight(_get_field(item, F_TITLE)),
            "type": _clean_highlight(_get_field(item, F_TYPE)),
            "date": _get_field(item, F_DATE).split(" ")[0] if _get_field(item, F_DATE) else "",
            "org": _get_field(item, "f_20232151076") or _get_field(item, F_ORG) or _get_field(item, F_ORG2),
            "doc_num": _get_field(item, F_DOC_NUM),
            "content": _get_field(item, F_CONTENT),
            "url": _get_field(item, F_URL),
        })
    return {"total": total, "rows": parsed}


def download_as_docx(keyword: str, output_dir: str = ".", reg_type: str = None) -> str:
    """搜索规章并将详情页 HTML 转为 docx 文件（通过 pandoc）"""
    # Step 1: search to find the doc_pub_url
    resp = search_with_content(keyword, reg_type=reg_type, size=5)
    rows = resp.get("rows", [])
    clean = re.sub(r'[《》<>〈〉]', '', keyword).strip()
    best = None
    for row in rows:
        title = row.get("title", "")
        title_clean = re.sub(r'[《》<>〈〉]', '', title).strip()
        if title_clean == clean:
            if best is None or row.get("type") == "部门规章":
                best = row
    if best is None:
        for row in rows:
            title_clean = re.sub(r'[《》<>〈〉]', '', row.get("title", "")).strip()
            if clean in title_clean and len(clean) > len(title_clean) * 0.5:
                if best is None or row.get("type") == "部门规章":
                    best = row
    if not best:
        raise ValueError(f"未找到 '{keyword}'")

    title = best["title"]
    date = best.get("date", "")
    typ = best.get("type", "规章")
    url = best.get("url", "")

    fname = f"{typ}_{title}_{date}_有效.docx" if date else f"{typ}_{title}_有效.docx"
    fname = re.sub(r'[\\/:*?"<>|]', '_', fname)
    fpath = os.path.join(output_dir, fname)
    os.makedirs(output_dir, exist_ok=True)

    # Step 2: try to fetch detail page HTML and convert via pandoc
    if url:
        try:
            return _download_via_html(url, title, fpath)
        except Exception:
            pass  # fallback to content-based approach

    # Step 3: fallback - generate docx from search result content
    content = best.get("content", "")
    if not content:
        raise ValueError(f"未找到 '{keyword}' 的全文内容且无法访问详情页")
    return _generate_docx_from_text(content, title, fpath)


def _download_via_html(url: str, title: str, fpath: str) -> str:
    """Fetch detail page HTML, extract article body, convert to docx via pandoc."""
    import subprocess, tempfile

    req = urllib.request.Request(url, headers={
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7)"
    })
    with urllib.request.urlopen(req, timeout=30) as resp:
        html = resp.read().decode("utf-8", errors="replace")

    # Extract article body from pages_content div
    body = _extract_article_body(html)
    if not body or len(body) < 100:
        raise ValueError("Failed to extract article body from HTML")

    # Wrap in standalone HTML
    standalone = f'<html><head><meta charset="utf-8"><title>{title}</title></head><body>\n{body}\n</body></html>'

    # Write temp HTML and convert with pandoc
    tmp_html = os.path.join(tempfile.gettempdir(), "gov_article_tmp.html")
    with open(tmp_html, "w", encoding="utf-8") as f:
        f.write(standalone)

    result = subprocess.run(
        ["pandoc", tmp_html, "-f", "html", "-t", "docx", "-o", fpath],
        capture_output=True, text=True, timeout=30
    )
    if result.returncode != 0:
        raise ValueError(f"pandoc failed: {result.stderr}")

    return fpath


def _extract_article_body(html: str) -> str:
    """Extract the article content div from a gov.cn page."""
    start = html.find('class="pages_content"')
    if start < 0:
        start = html.find('class="trs_editor_view')
    if start < 0:
        return ""

    pos = html.find(">", start) + 1
    depth = 1
    while depth > 0 and pos < len(html):
        next_open = html.find("<div", pos)
        next_close = html.find("</div>", pos)
        if next_close < 0:
            break
        if next_open >= 0 and next_open < next_close:
            depth += 1
            pos = next_open + 4
        else:
            depth -= 1
            if depth == 0:
                return html[html.find(">", start) + 1 : next_close]
            pos = next_close + 6
    return ""


def _generate_docx_from_text(content: str, title: str, fpath: str) -> str:
    """Fallback: generate docx from plain text content."""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading(title, level=1)
        text = content.strip()
        segments = re.split(r'(?=第[一二三四五六七八九十百千\d]+[条章节编])', text)
        for seg in segments:
            seg = seg.strip()
            if not seg:
                continue
            sub_parts = re.split(r'(?=（[一二三四五六七八九十\d]+）)', seg)
            for part in sub_parts:
                part = part.strip()
                if part:
                    p = doc.add_paragraph(part)
                    if re.match(r'^第[一二三四五六七八九十百千\d]+[条章节编]', part):
                        for run in p.runs:
                            run.bold = True
        doc.save(fpath)
    except ImportError:
        # No python-docx, save as txt
        txt_path = fpath.replace(".docx", ".txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(content)
        return txt_path
    return fpath


# ── CLI ──

def cmd_search(args):
    if args.all_pages:
        resp = search_all_pages(args.keyword, reg_type=args.type,
                                content_search=args.content, page_size=args.size)
    else:
        resp = search(args.keyword, reg_type=args.type,
                      content_search=args.content, size=args.size)
    print(f"共 {resp.get('total', 0)} 条结果（返回 {len(resp.get('rows',[]))} 条）：")
    for r in resp.get("rows", []):
        print(f"  [{r.get('type','')}] {r['title']} | {r.get('date','')} | {r.get('org','')}")
        if r.get("url"):
            print(f"    → {r['url']}")


def cmd_batch_check(args):
    with open(args.file, "r", encoding="utf-8") as f:
        names = [line.strip() for line in f if line.strip()]
    print(f"共 {len(names)} 条规章待检查", file=sys.stderr)
    results = batch_check(names, delay=args.delay)
    if args.json:
        print(json.dumps(results, ensure_ascii=False, indent=2))
    else:
        for r in results:
            mark = "✓" if r["found"] else "✗"
            print(f"  {mark} {r['name']} → {r.get('status','?')} ({r.get('org','')})")
        found = sum(1 for r in results if r["found"])
        not_found = sum(1 for r in results if not r["found"])
        print(f"\n汇总: 找到 {found}, 未找到 {not_found}")


def main():
    parser = argparse.ArgumentParser(description="国家规章库查询工具")
    sub = parser.add_subparsers(dest="command", required=True)

    p = sub.add_parser("search", help="搜索规章")
    p.add_argument("keyword", help="搜索关键词")
    p.add_argument("--type", help="规章类型: 部门规章/地方政府规章")
    p.add_argument("--content", action="store_true", help="按全文内容搜索（默认按标题）")
    p.add_argument("--all-pages", action="store_true", help="获取所有页结果")
    p.add_argument("--size", type=int, default=10, help="每页条数")
    p.set_defaults(func=cmd_search)

    p = sub.add_parser("batch-check", help="批量检查规章状态")
    p.add_argument("file", help="规章名称列表文件")
    p.add_argument("--delay", type=float, default=0.5, help="请求间隔秒数")
    p.add_argument("--json", action="store_true", help="输出 JSON")
    p.set_defaults(func=cmd_batch_check)

    args = parser.parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
